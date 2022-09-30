from __future__ import print_function
import sys
import datetime
import pickle
import os.path
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
import speech_recognition as sr
import webbrowser
import time
import playsound
import os
import random
import requests
import sys
import ssl
import certifi
import docx
from reportlab.pdfgen.canvas import Canvas
from gtts import gTTS
from time import ctime
from time import strftime
from PyQt5 import QtCore, QtGui
from PyQt5.QtCore import Qt, QUrl
from PyQt5.QtGui import QIcon, QPalette
from PyQt5.QtMultimedia import QMediaContent, QMediaPlayer
from PyQt5.QtMultimediaWidgets import QVideoWidget
from PyQt5.QtWidgets import (QApplication, QDialog, QFileDialog, QGroupBox,
                             QHBoxLayout, QLabel, QMainWindow, QPushButton,
                             QRadioButton, QSizePolicy, QSlider, QStyle,
                             QVBoxLayout, QWidget, QMessageBox)

r = sr.Recognizer()


def record_audio(ask=False):
    with sr.Microphone() as source:
        if ask:
            Assistant_speak(ask)
        audio = r.listen(source)
        voice_data = ''
        try:
            voice_data = r.recognize_google(audio)
        except sr.UnknownValueError:
            Assistant_speak('Sorry, I did not get that')
        except sr.RequestError:
            Assistant_speak('Sorry, my speech service is down')
        return voice_data


def Assistant_speak(audio_string):
    tts = gTTS(text=audio_string, lang='en')
    r = random.randint(1, 10000000)
    audio_file = 'audio-' + str(r) + '.mp3'
    tts.save(audio_file)
    playsound.playsound(audio_file)
    print(audio_string)
    os.remove(audio_file)

class MainWindow(QWidget):
    def __init__(self):
        super().__init__()

        self.title = "Voice Assistant"
        self.top = 100
        self.left = 100
        self.width = 400
        self.height = 600

        self.recWindow = QWidget() # separate record window for when record button pressed
        self.impWindow = QWidget() # import window
        self.liveWindow = QWidget() # live window

        self.initWindow()

    def initWindow(self):
        self.setWindowIcon(QtGui.QIcon("MainIcon.png"))
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)

        self.createLayout()
        vbox = QVBoxLayout()
        vbox.addWidget(self.groupBox)
        vbox.addWidget(self.genBox)
        vbox.addWidget(self.webBox)
        vbox.addWidget(self.exitButton)
        self.setLayout(vbox)

        self.show()

    def createLayout(self):
        # transcribe section
        self.groupBox = QGroupBox("Transcribe")
        hBoxLayout = QVBoxLayout()

        impButton = QPushButton("Import", self)
        hBoxLayout.addWidget(impButton)
        impButton.setToolTip("Click to import an audio file")
        impButton.clicked.connect(self.importClicked)
        
        liveButton = QPushButton("Record", self)
        hBoxLayout.addWidget(liveButton)
        liveButton.setToolTip("Click to record your words and save it as text and audio file")
        liveButton.clicked.connect(self.liveClicked)

        self.groupBox.setLayout(hBoxLayout)

        # general assistant section
        self.genBox = QGroupBox("General assistant")
        gBoxLayout = QVBoxLayout()

        generalAssistantButton = QPushButton("Ask something", self)
        gBoxLayout.addWidget(generalAssistantButton)
        generalAssistantButton.setToolTip("Click to ask the assistant something")
        generalAssistantButton.clicked.connect(self.generalAssistantClicked)

        self.genBox.setLayout(gBoxLayout)

        # web assistant section
        self.webBox = QGroupBox("Web assistant")
        wBoxLayout = QVBoxLayout()

        # google search
        gSearchButton = QPushButton("Google Search", self)
        wBoxLayout.addWidget(gSearchButton)
        gSearchButton.setToolTip("Click to search Google with your voice")
        gSearchButton.clicked.connect(self.googleSearchClicked)

        # find location
        locationButton = QPushButton("Find a location", self)
        wBoxLayout.addWidget(locationButton)
        locationButton.setToolTip("Click to find a location on Google Maps")
        locationButton.clicked.connect(self.findLocationClicked)

        # play music
        playMusicButton = QPushButton("Youtube", self)
        wBoxLayout.addWidget(playMusicButton)
        playMusicButton.setToolTip("Click to play music from Spotify")
        playMusicButton.clicked.connect(self.playMusicClicked)

        self.webBox.setLayout(wBoxLayout)

        # exit section
        self.exitButton = QPushButton("Exit", self)
        self.exitButton.setToolTip("Click to exit the application")
        self.exitButton.setStyleSheet("background: darkred; color: white")
        self.exitButton.clicked.connect(self.exitClicked)

    def importClicked(self):
        # code for import button
        self.impWindow = ImportWindow() # open import window (ImportWindow class)

    def liveClicked(self):
            def speak(audio_string):
                tts = gTTS(text=audio_string, lang='en') # text to speech(voice)
                r = random.randint(1,20000)
                audio_file = 'Transcribe' + str(r) + '.wav'
                tts.save(audio_file) # save as wav
                print(f"Transcribe: {audio_string}") # print what app said

            ''' recording the sound '''
            with sr.Microphone() as source:
                Assistant_speak("Adjusting noise ")
                r.adjust_for_ambient_noise(source, duration=1)
                Assistant_speak("Recording is starting ")
                audio = r.listen(source)
                voice_data = ''
                Assistant_speak("Done recording")
            
            ''' Recorgnizing the Audio '''
            try:
                    voice_data = r.recognize_google(audio)  # convert audio to text
            except sr.UnknownValueError: # error: recognizer does not understand
                    print('I did not get that')
            except sr.RequestError:
                    print('Sorry, the service is down') # error: recognizer is not connected
            speak(f"{voice_data.lower()}") # print what user said
        
            with open("transcribe.txt", "w") as text_file:
                text_file.write(voice_data.lower())
                

    def generalAssistantClicked(self):
        # code to activate general assistant
        def respond(voice_data):
            # pre-programmed responses
            if 'name' in voice_data:
                Assistant_speak('My name is Red')
            elif 'what are you' in voice_data:
                Assistant_speak('I am your personal voice assistant. How may i help?')
            elif 'heads or tails' in voice_data:
                Assistant_speak('Huh, it landed on the side')
            elif 'date' in voice_data:
                Assistant_speak(ctime())
            ##elif 'reminders' in voice_data:
            ##  service = authenticate_google()
            ##service = get_events(5, service)
            elif 'joke' in voice_data:
                Assistant_speak('knock knock... oh wait how did it go again')
            elif 'story' in voice_data:
                Assistant_speak('A long time ago, in a galaxy far, far away...')
            elif 'who made you' in voice_data or 'who built you' in voice_data or 'who created you' in voice_data:
                Assistant_speak('I was built by students')
                print('I was built by some students.')
            elif 'features' in voice_data:
                Assistant_speak('Here is a list of my features')
                print(
                    'Google search, Location finder, Youtube, Weather, Event manager, time/date, flip a coin, tell a joke ')
            elif 'hello' in voice_data:
                day_time = int(strftime('%H'))
                if day_time < 12:
                    Assistant_speak('Hello, Good morning')
                elif 12 <= day_time < 18:
                    Assistant_speak('Hello, Good afternoon')
                else:
                    Assistant_speak('Hello, Good evening')
            elif 'search' in voice_data:
            #Say search followed by what you want to search
                Assistant_speak('What would you like to search?')
                search = record_audio('')
                url = 'https://google.com/search?q=' + search
                webbrowser.get().open(url)
                time.sleep(3)

            elif 'Maps' in voice_data:
                #in maps search" followed by location
                location = record_audio('')
                url = 'https://google.nl/maps/place/' + location + '/&amp;'
                webbrowser.get().open(url)
                Assistant_speak('lets go ' + location)
                time.sleep(3)

            elif 'Youtube' in voice_data:
                url = 'https://www.youtube.com/'
                webbrowser.get().open(url)
                Assistant_speak('Opening Youtube')
                time.sleep(3)

            if "weather" in voice_data:
                api_key = "8ef61edcf1c576d65d836254e11ea420"
                base_url = "https://api.openweathermap.org/data/2.5/weather?"
                Assistant_speak("whats the city name")
                city_name = record_audio()
                complete_url = base_url + "appid=" + api_key + "&q=" + city_name
                response = requests.get(complete_url)
                x = response.json()
                if x["cod"] != "404":
                    y = x["main"]
                    current_temperature = y["temp"]
                    z = x["weather"]
                    weather_description = z[0]["description"]
                    Assistant_speak(" Temperature in kelvin unit is " +
                                    str(current_temperature) +
                                    "\n description  " +
                                    str(weather_description))
                    print(" Temperature in kelvin unit = " +
                          str(current_temperature) +
                          "\n description = " +
                          str(weather_description))
                else:
                    Assistant_speak(" City Not Found ")


            if 'exit' in voice_data or 'quit' in voice_data:
                Assistant_speak('Good bye, Have a nice day')
                exit()
        time.sleep(1)
        Assistant_speak("How can I help?")
        while 1:
            voice_data = record_audio()
            respond(voice_data)



    def googleSearchClicked(self):
        # code to search on google
        # Say search followed by what you want to search
        Assistant_speak("What do you want to search on Google?")
        search = record_audio('')
        url = 'https://google.com/search?q=' + search
        webbrowser.get().open(url)

    def findLocationClicked(self):
        # code to find a location
        # in maps search" followed by location
        Assistant_speak("What location do you want to search on Google?")
        location = record_audio('')
        url = 'https://google.nl/maps/place/' + location + '/&amp;'
        webbrowser.get().open(url)
        Assistant_speak('lets go:' + location)

    def playMusicClicked(self):
        # code to play music
        url = 'https://www.youtube.com'
        webbrowser.get().open(url)
        Assistant_speak('Opening Youtube')
        time.sleep(3)

    def exitClicked(self):
        sys.exit(App.exec()) # exit the application

class ImportWindow(QWidget):
    def __init__(self):
        self.path = ""
        self.ext= ""
        super().__init__()

        self.title = "Import"
        self.top = 100
        self.left = 100
        self.width = 400
        self.height = 200

        self.transcribeButton = QPushButton()
        self.startImpButton = QPushButton()

        self.initWindow()

    def initWindow(self):
        self.setWindowIcon(QtGui.QIcon("ImportIcon.png"))
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)

        self.createLayout()
        vbox = QVBoxLayout()
        vbox.addWidget(self.radioBox)
        vbox.addWidget(self.groupBox)

        self.setLayout(vbox)

        self.show()

    def createLayout(self):
        self.groupBox = QGroupBox("Importing")
        hBoxLayout = QVBoxLayout()

        self.startImpButton = QPushButton("IMPORT", self)
        hBoxLayout.addWidget(self.startImpButton)
        self.startImpButton.setToolTip("Click to import a sound file")
        self.startImpButton.clicked.connect(self.imp)

# radio buttons
        self.radioBox = QGroupBox("Choose file type")
        rBoxLayout = QVBoxLayout()

        self.r0 = QRadioButton("docx")
        self.r0.setText(".docx")
        self.r0.toggled.connect(lambda: self.radioState(self.r0))
        self.r0.clicked.connect(self.impTranscribe)

        self.r1 = QRadioButton("pdf")
        self.r1.setText(".pdf")
        self.r1.toggled.connect(lambda: self.radioState(self.r1))
        self.r1.clicked.connect(self.impTranscribe)

        self.r2 = QRadioButton("txt")
        self.r2.setText(".txt")
        self.r2.toggled.connect(lambda: self.radioState(self.r2))
        self.r2.clicked.connect(self.impTranscribe)

        rBoxLayout.addWidget(self.r0)
        rBoxLayout.addWidget(self.r1)
        rBoxLayout.addWidget(self.r2)
        self.r0.setEnabled(False)
        self.r1.setEnabled(False)
        self.r2.setEnabled(False)

        self.radioBox.setLayout(rBoxLayout)

        self.transcribeButton = QPushButton("TRANSCRIBE", self)
        hBoxLayout.addWidget(self.transcribeButton)
        self.transcribeButton.setToolTip(
            "Click to transcribe an imported sound file")
        self.transcribeButton.clicked.connect(self.transcribe)
        # user cannot click this button until they have imported something
        self.transcribeButton.setEnabled(False)

        self.groupBox.setLayout(hBoxLayout)

    def radioState(self, b):
        if b.isChecked() == True:
            self.ext = b.text()
            print(self.ext)
            print(b.text()+" is selected")

        else:
            print(self.ext)
            print(b.text()+" is deselected")
            

    def impTranscribe(self):

        typeChosen = True
        if(typeChosen):
            self.transcribeButton.setEnabled(True)

    def imp(self):
        # code to import
        filename = QFileDialog.getOpenFileName()
        filePath = filename[0]
        self.path = os.path.basename(os.path.normpath(filePath))
        print(self.path)
        #print(filename)
        importSuccess = True

        # code here to import and check its successful
        # if not, set importSuccess to false

        if(importSuccess):
            self.r0.setEnabled(True)
            self.r1.setEnabled(True)
            self.r2.setEnabled(True)

    def transcribe(self):
        # code to transcribe imported file
        print(self.path)
        print(self.ext)

        #splits filename so new extension can be added
        fileName = self.path.split('.')[0]
        fileFormat = self.ext
        sound = self.path

        r = sr.Recognizer()

        
        with sr.AudioFile(sound) as source:
            r.adjust_for_ambient_noise(source)

            print("Transcribing Audio File...")

            audio = r.listen(source)

            #r.recognize_google(audio) this is the Tranbscribed audio to be written to file
            try:
                print("Converted Audio Is : \n " + r.recognize_google(audio))

                #formats using differnet code snippets depending on file format
                if fileFormat == ".docx":# Write to docx
                    print("Audio is being Converted to docx format....")
                    mydoc = docx.Document()
                    mydoc.add_heading(fileName, 0)
                    mydoc.add_paragraph(r.recognize_google(audio))
                    mydoc.save("./recordings/" + fileName + fileFormat )
                    print("conversion complete")
                    QMessageBox.information(self, "Success", "File successfully Transcribed")
                    self.close()

                elif fileFormat == ".pdf":  #Write to text
                    print("Audio is being Converted to PDF format....")
                    canvas = Canvas("./recordings/" + fileName + fileFormat)
                    canvas.setTitle(fileName)
                    canvas.drawString(72,800, fileName)
                    canvas.drawString(72, 700, r.recognize_google(audio))
                    canvas.save()
                    print("conversion complete")
                    QMessageBox.information(self, "Success", "File successfully Transcribed")
                    self.close()

                else:# Write to text
                    print("Audio is being Converted to Text format....")
                    f = open("./recordings/" + fileName + fileFormat, "x")
                    f.write(r.recognize_google(audio))
                    f.close()
                    print("conversion complete")

                    self.close()
                    QMessageBox.information(self, "Success", "File successfully Transcribed")

            except Exception as e:
                print(e)
                #converts error to string for GUI Display str(e)
                QMessageBox.critical(self, "Success", str(e))

if __name__ == "__main__":
    App = QApplication(sys.argv)
    window = MainWindow()
    sys.exit(App.exec())